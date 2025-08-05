from pathlib import Path
from typing import List, Dict, Any, Optional
import logging
import mimetypes
from fastapi import UploadFile
import PyPDF2
import docx
from PIL import Image
import pytesseract

from llama_index.core import Document
from llama_index.readers.file import (
    PyMuPDFReader,
    UnstructuredReader,
    PagedCSVReader
)

logger = logging.getLogger(__name__)

ALLOWED_MEDICAL_EXTENSIONS = {'.pdf', '.docx', '.txt', '.jpg', '.jpeg', '.png', '.tiff'}
ALLOWED_MIME_TYPES = {
    'application/pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'text/plain',
    'image/jpeg',
    'image/png',
    'image/tiff'
}

def validate_medical_file(file: UploadFile) -> bool:
    """Validate if uploaded file is a supported medical document type"""
    allowed_types = {
        'application/pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'text/plain',
        'image/jpeg',
        'image/png'
    }
    
    allowed_extensions = {'.pdf', '.docx', '.txt', '.jpg', '.jpeg', '.png'}
    
    # Check content type
    if file.content_type not in allowed_types:
        return False
    
    # Check file extension
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in allowed_extensions:
        return False
    
    return True

def extract_text_from_file(file_path: str, content_type: str) -> str:
    """Extract text content from various file types"""
    try:
        if content_type == 'application/pdf':
            # Implement PDF text extraction
            import PyPDF2
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                return text
        
        elif content_type == 'text/plain':
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        
        elif content_type in ['image/jpeg', 'image/png']:
            # Implement OCR for images
            # You'll need to install pytesseract and Pillow
            import pytesseract
            from PIL import Image
            image = Image.open(file_path)
            return pytesseract.image_to_string(image)
        
        elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            # Implement DOCX text extraction
            from docx import Document
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        
        else:
            raise ValueError(f"Unsupported file type: {content_type}")
            
    except Exception as e:
        raise ValueError(f"Error extracting text from file: {str(e)}")

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file."""
    text = ""
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file."""
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def extract_text_from_image(file_path: str) -> str:
    """Extract text from image using OCR."""
    image = Image.open(file_path)
    text = pytesseract.image_to_string(image)
    return text

class DocumentProcessor:
    """Utility class for processing different document types"""
    
    def __init__(self):
        self.pdf_reader = PyMuPDFReader()
        self.unstructured_reader = UnstructuredReader()
        self.csv_reader = PagedCSVReader()
    
    def process_pdf(self, file_path: Path) -> List[Document]:
        """Process PDF files"""
        try:
            return self.pdf_reader.load_data(file_path=file_path, metadata=True)
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            return []
    
    def process_unstructured(self, file_path: Path) -> List[Document]:
        """Process various file types using Unstructured"""
        try:
            return self.unstructured_reader.load_data(
                unstructured_kwargs={"filename": str(file_path)}
            )
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            return []
    
    def process_csv(self, file_path: Path) -> List[Document]:
        """Process CSV files"""
        try:
            return self.csv_reader.load_data(file=file_path)
        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {e}")
            return []
    
    def process_directory(self, directory_path: Path) -> List[Document]:
        """Process all supported files in a directory"""
        documents = []
        
        for file_path in directory_path.rglob("*"):
            if file_path.is_file():
                suffix = file_path.suffix.lower()
                
                if suffix == ".pdf":
                    documents.extend(self.process_pdf(file_path))
                elif suffix == ".csv":
                    documents.extend(self.process_csv(file_path))
                elif suffix in [".txt", ".md", ".html", ".docx"]:
                    documents.extend(self.process_unstructured(file_path))
        
        return documents